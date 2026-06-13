from flask_cors import CORS
from flask import Flask, request, jsonify, session
from flask_cors import CORS
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
import docx, requests, os, uuid
import mammoth
import re
import sqlite3
from werkzeug.security import generate_password_hash, check_password_hash

# -------------------------------------------------------
# APP CONFIG
# -------------------------------------------------------
app = Flask(__name__)
app.secret_key = "your_secret_key_here"
CORS(app, resources={r"/api/*": {"origins": "*"}}, supports_credentials=True)

UPLOAD_FOLDER = "uploads"
ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

HF_API_KEY = os.getenv("HF_API_KEY")
HEADERS = {"Authorization": f"Bearer {HF_API_KEY}", "Content-Type": "application/json"}

# -------------------------------------------------------
# HELPERS
# -------------------------------------------------------
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(path):
    ext = path.rsplit(".", 1)[1].lower()
    
    # DOCX via Mammoth
    if ext == "docx":
        try:
            with open(path, "rb") as f:
                text = mammoth.extract_raw_text(f).value.strip()
                return text
        except:
            try:
                d = docx.Document(path)
                return "\n".join(p.text for p in d.paragraphs)
            except:
                return ""
    
    # PDF
    elif ext == "pdf":
        try:
            reader = PdfReader(path)
            pages = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages.append(t.replace("\n", " "))
            return "\n".join(pages)
        except:
            return ""
    
    # TXT
    elif ext == "txt":
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except:
            return ""
    
    return ""

def clean_text(text):
    text = text.replace("\r", " ").replace("\t", " ")
    text = " ".join(text.split())
    return text

def enforce_word_limit(text, target_words):
    words = text.split()
    return " ".join(words[:target_words])


# -------------------------------------------------------
# API: SUMMARIZATION (PEGASUS 2-PASS FOR 750 WORDS)
# -------------------------------------------------------
@app.route("/api/upload", methods=["POST"])
def upload_and_summarize():

    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    length = int(request.form.get("length", "250"))

    if not allowed_file(file.filename):
        return jsonify({"error": "Unsupported file type"}), 400

    # Save incoming file
    filename = secure_filename(file.filename)
    unique_name = f"{uuid.uuid4().hex}_{filename}"
    path = os.path.join(UPLOAD_FOLDER, unique_name)
    file.save(path)

    # Extract text
    text = extract_text(path)
    if not text.strip():
        return jsonify({"error": "Could not extract text"}), 400

    # Use only first 5000 chars for stability
    input_block = text[:5000]

    PEGASUS = "https://router.huggingface.co/hf-inference/models/google/pegasus-cnn_dailymail"

    def summarize_with_pegasus(content, target_words):
        payload = {
            "inputs": content,
            "parameters": {
                "max_length": target_words + 150, 
                "min_length": int(target_words*0.75),
                "do_sample": False,
                "num_beams": 6,
                "repetition_penalty": 1.7,
                "no_repeat_ngram_size": 4
            }
        }

        r = requests.post(PEGASUS, headers=HEADERS, json=payload)
        if r.status_code != 200:
            raise ValueError("Pegasus summarization failed: " + r.text)

        data = r.json()
        raw = data[0]["summary_text"].strip()
        return clean_summary(raw)

    # -----------------------------
    # CLEANING FUNCTION
    # -----------------------------
    import re
    def clean_summary(text):
        text = re.sub(r"http\S+", "", text)               # remove URLs
        text = re.sub(r"www\.\S+", "", text)             # remove www links
        text = re.sub(r"\S+@\S+", "", text)              # remove emails
        text = re.sub(r"\d{3}-\d{3}-\d{4}", "", text)    # US phone
        text = re.sub(r"\d{10}", "", text)               # numbers
        text = re.sub(r"for more information.*", "", text, flags=re.I)
        text = re.sub(r"visit .*", "", text, flags=re.I)
        text = re.sub(r"this essay .*", "", text, flags=re.I)
        text = re.sub(r"watch .*", "", text, flags=re.I)
        text = re.sub(r"order .*", "", text, flags=re.I)
        text = re.sub(r"buy .*", "", text, flags=re.I)
        text = " ".join(text.split())                   # normalize spaces
        return text.strip()

    # ---------------------------------------------
    # 250 & 500-word summaries (direct Pegasus)
    # ---------------------------------------------
    if length <= 500:
        summary = summarize_with_pegasus(input_block, length)
        summary = enforce_word_limit(summary, length)
        return jsonify({"summary": summary})

    # ---------------------------------------------
    # 750-word summary (2-pass Pegasus)
    # ---------------------------------------------
    # Pass 1 → 400 words
    half1 = summarize_with_pegasus(input_block, 400)

    # Pass 2 → refine half1 into another 400 words
    half2 = summarize_with_pegasus(half1, 400)

    # Merge
    merged = half1 + " " + half2
    merged = clean_summary(merged)

    final_summary = enforce_word_limit(merged, 750)

    return jsonify({"summary": final_summary})

#-----------------------------------------------------------------------------------------------------------
#--------------------NOTES GENERATION ----------------------------------------------------------------------
#-----------------------------------------------------------------------------------------------------------
@app.route("/api/notes", methods=["POST"])
def generate_notes():
    """
    FINAL NOTES GENERATOR
    • Large files → 8–9 bullets, each 70–90 words
    • Small files → 5–8 bullets, each 40–55 words
    • Always full sentences
    • Never starts or ends mid-sentence
    """

    import re
    import time

    try:
        # ---------------------------------------------------
        # 1. Receive File
        # ---------------------------------------------------
        file = request.files.get("file")
        if not file:
            return jsonify({"error": "No file uploaded"}), 400

        filename = secure_filename(file.filename)
        save_path = os.path.join("uploads", filename)
        os.makedirs("uploads", exist_ok=True)
        file.save(save_path)

        # ---------------------------------------------------
        # 2. Extract Text
        #---------------------------------------------------
        text = ""
        ext = filename.lower().split(".")[-1]

        if ext == "docx":
            try:
                with open(save_path, "rb") as f:
                    text = mammoth.extract_raw_text(f).value.strip()
            except:
                text = ""

            if not text.strip():
                try:
                    d = docx.Document(save_path)
                    text = "\n".join(p.text for p in d.paragraphs)
                except:
                    text = ""

        elif ext == "pdf":
            try:
                reader = PdfReader(save_path)
                pages = []
                for pg in reader.pages:
                    t = pg.extract_text()
                    if t:
                        pages.append(t.replace("\n", " "))
                text = "\n".join(pages)
            except:
                text = ""

        elif ext == "txt":
            try:
                with open(save_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            except:
                text = ""

        if not text.strip():
            return jsonify({"error": "Could not extract text"}), 400

        # ---------------------------------------------------
        # 3. Clean Text
        # ---------------------------------------------------
        text = text.replace("\r", " ").replace("\t", " ")
        text = " ".join(text.split())
        words = text.split()
        wc = len(words)

        # ---------------------------------------------------
        # 4. Decide Small / Large File Logic
        # ---------------------------------------------------
        if wc < 1000:
            model = "pegasus"
            bullet_min, bullet_max = 40, 55
            total_notes = min(8, max(5, wc // 150))
        else:
            model = "bart"
            bullet_min, bullet_max = 70, 90
            total_notes = 9

        PEGASUS_URL = "https://api-inference.huggingface.co/models/google/pegasus-xsum"
        BART_URL = "https://api-inference.huggingface.co/models/facebook/bart-large-cnn"
        headers = {"Authorization": f"Bearer {HF_API_KEY}"}

        # ---------------------------------------------------
        # 5. Sentence-Aligned Chunking (Fix)
        # ---------------------------------------------------
        def sentence_aligned_chunks(text, target_words=350):
            sentences = re.findall(r'[^.!?]*[.!?]', text)
            chunks, current = [], []
            wc = 0

            for s in sentences:
                w = len(s.split())
                if wc + w > target_words:
                    chunks.append(" ".join(current).strip())
                    current = []
                    wc = 0
                current.append(s.strip())
                wc += w

            if current:
                chunks.append(" ".join(current).strip())

            return chunks

        chunks = sentence_aligned_chunks(text, target_words=350)

        # ---------------------------------------------------
        # 6. Trim to 70–90 Word Full-Sentence Bullet
        # ---------------------------------------------------
        def trim_full_sentence_bullet(sentence_list, min_w, max_w):
            final, wc = [], 0

            for s in sentence_list:
                w = len(s.split())
                if wc + w > max_w:
                    break
                final.append(s)
                wc += w
                if wc >= min_w:
                    break

            if wc < min_w:
                combined = " ".join(sentence_list)
                fallback = " ".join(combined.split()[:max_w])
                return fallback.strip() + "."

            bullet = " ".join(final).strip()
            if not bullet.endswith((".", "?", "!")):
                bullet += "."
            return bullet

        # ---------------------------------------------------
        # 7. Model Bullet Generation
        # ---------------------------------------------------
        def generate_bullet(chunk):
            prompt = (
                f"Write ONE study bullet of about {bullet_min}-{bullet_max} words. "
                "Your output must be composed of full sentences ending with '.', '?', or '!'. "
                "Never cut a sentence midway. No filler, no examples.\n\n"
                f"TEXT:\n{chunk}"
            )

            try:
                if model == "pegasus":
                    resp = requests.post(PEGASUS_URL, headers=headers,
                        json={"inputs": prompt}, timeout=60)
                else:
                    payload = {
                        "inputs": prompt,
                        "parameters": {
                            "max_length": 350,
                            "min_length": 60,
                            "num_beams": 4,
                            "do_sample": False
                        }
                    }
                    resp = requests.post(BART_URL, headers=headers,
                        json=payload, timeout=60)

                j = resp.json()
                if isinstance(j, list) and j:
                    raw = j[0].get("summary_text") or j[0].get("generated_text") or ""
                else:
                    raw = ""
            except:
                raw = ""

            raw = raw.strip().lstrip("0123456789).•- ")
            raw = " ".join(raw.split())  # collapse spaces

            sentences = re.findall(r'[^.!?]*[.!?]', raw)
            if not sentences:
                fallback = " ".join(chunk.split()[:bullet_max])
                return fallback.strip() + "."

            bullet = trim_full_sentence_bullet(sentences, bullet_min, bullet_max)
            return bullet

        # ---------------------------------------------------
        # 8. Generate Final Notes
        # ---------------------------------------------------
        final_notes = []

        for chunk in chunks:
            if len(final_notes) >= total_notes:
                break
            bullet = generate_bullet(chunk)
            final_notes.append(f"• {bullet}")
            time.sleep(0.5)

        return jsonify({"notes": final_notes})

    except Exception as e:
        return jsonify({"error": str(e)}), 500



# QUIZ API (unchanged)
# -------------------------------------------------------
@app.route("/api/quiz", methods=["POST"])
def generate_quiz():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    filename = secure_filename(file.filename)
    path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(path)

    text = extract_text(path)
    sentences = [s.strip() for s in text.split(".") if len(s.split()) > 5]

    if len(sentences) < 5:
        return jsonify({"error": "Not enough text"}), 400

    quiz = []
    for i, sent in enumerate(sentences[:10]):
        quiz.append({
            "question": f"What is the main idea of: '{sent[:80]}...'",
            "options": ["Main concept", "Minor detail", "Opposite meaning", "Unrelated idea"],
            "answer": "Main concept",
            "explanation": f"Based on: {sent}"
        })

    return jsonify({"quiz": quiz})


# -------------------------------------------------------
# RUN APP
# -------------------------------------------------------
if __name__ == "__main__":
    app.run(debug=True)
